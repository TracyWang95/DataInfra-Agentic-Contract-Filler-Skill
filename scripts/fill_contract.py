# -*- coding: utf-8 -*-
"""
Fill the DOCX template with collected field values and generate the final contract.
Works with any contract type - reads config from state file.

Usage:
    python fill_contract.py --state "path/to/state.json" --output "path/to/output.docx"
    
    # Optionally override template:
    python fill_contract.py --state "state.json" --template "custom.docx" --output "out.docx"
"""

import argparse
import json
import re
import sys
from pathlib import Path

# Fix Windows console encoding
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

SKILL_DIR = Path(__file__).parent.parent
sys.path.insert(0, str(SKILL_DIR))

from contracts.base_config import (
    apply_aliases,
    get_config,
    get_unfilled_fields,
    is_checkbox_checked,
    is_field_filled,
)


# Style constants matching templates
FONT_BODY = "仿宋"
SZ_BODY = 14  # in points


def fill_docx_template(template_path: str, field_values: dict, output_path: str) -> list[str]:
    """
    Replace {{placeholder}} markers in a DOCX template with actual values.
    
    Returns:
        list of unfilled placeholder names
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn
    except ImportError:
        print("ERROR: python-docx is required. Install with: pip install python-docx")
        sys.exit(1)
    
    doc = Document(template_path)
    unfilled = set()
    filled_count = 0
    
    def replace_in_text(text: str, values: dict) -> tuple[str, set, int]:
        """Replace all {{key}} in text with values[key]."""
        remaining = set()
        count = 0
        
        def replacer(m):
            nonlocal count
            key = m.group(1)
            
            # 复选框字段（以 ☐ 开头）
            if key.startswith("☐"):
                if key in values:
                    val = values[key]
                    if is_field_filled(key, values):
                        count += 1
                        return "☑" if is_checkbox_checked(val) else "☐"
                    # Invalid checkbox value should be treated as unfilled
                    remaining.add(key)
                    return "☐"
                else:
                    # 未填写的复选框
                    remaining.add(key)
                    return "☐"
            
            # 普通文本字段
            elif key in values and is_field_filled(key, values):
                count += 1
                return str(values[key]).strip()
            else:
                remaining.add(key)
                return ""  # clear unfilled placeholders
        
        result = re.sub(r'\{\{(.+?)\}\}', replacer, text)
        return result, remaining, count
    
    def process_paragraph(para, values: dict) -> int:
        """Process a single paragraph, preserving formatting."""
        full_text = para.text
        if '{{' not in full_text:
            return 0
        
        new_text, remaining, count = replace_in_text(full_text, values)
        unfilled.update(remaining)
        
        if new_text != full_text:
            # Preserve formatting from the first run
            if para.runs:
                font_name = para.runs[0].font.name or FONT_BODY
                font_size = para.runs[0].font.size
                font_bold = para.runs[0].font.bold
            else:
                font_name = FONT_BODY
                font_size = Pt(SZ_BODY)
                font_bold = False
            
            # Clear all runs and rewrite with preserved formatting
            for run in para.runs:
                run.text = ""
            
            if para.runs:
                para.runs[0].text = new_text
            else:
                r = para.add_run(new_text)
                r.font.name = font_name
                r._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                if font_size:
                    r.font.size = font_size
                if font_bold is not None:
                    r.font.bold = font_bold
        
        return count
    
    # Process paragraphs
    for para in doc.paragraphs:
        filled_count += process_paragraph(para, field_values)
    
    # Process tables (including nested tables) — same formatting logic as paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    filled_count += process_paragraph(para, field_values)
    
    # Save
    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    
    return sorted(unfilled)


def main():
    parser = argparse.ArgumentParser(description="填充合同模板生成成品")
    parser.add_argument("--state", required=True, help="状态文件路径（JSON）")
    parser.add_argument("--template", help="模板文件路径（可选，默认从状态文件读取）")
    parser.add_argument("--output", required=True, help="输出文件路径（.docx）")
    parser.add_argument("--force", action="store_true", help="强制生成，即使有未填字段")
    parser.add_argument("--check", action="store_true", help="仅检查未填字段，不生成文件")
    args = parser.parse_args()
    
    # Load state
    state_path = Path(args.state)
    if not state_path.exists():
        print(f"ERROR: 状态文件不存在：{state_path}")
        sys.exit(1)
    
    with open(state_path, "r", encoding="utf-8") as f:
        state = json.load(f)
    
    # Get template path
    template_path = args.template or state.get("template_path")
    if not template_path or not Path(template_path).exists():
        print(f"ERROR: 模板文件不存在：{template_path}")
        sys.exit(1)
    
    # Get field values and apply aliases
    field_values = state.get("field_values", {})
    
    # Load config for this contract type to get aliases and placeholder groups
    contract_type = state.get("contract_type")
    placeholder_groups = state.get("groups", {})
    
    if contract_type:
        try:
            config = get_config(contract_type)
            field_aliases = config.get("FIELD_ALIASES", {})
            field_values = apply_aliases(field_values, field_aliases)
            placeholder_groups = config.get("PLACEHOLDER_GROUPS", placeholder_groups)
        except Exception:
            pass  # If config can't be loaded, proceed without aliases
    
    # Show progress
    contract_name = state.get("contract_name", "合同")
    total_fields = state.get("total_placeholders", 0)
    filled_count = sum(1 for f in state.get("all_placeholders", []) if is_field_filled(f, field_values))
    
    print(f"📋 合同类型：{contract_name}")
    print(f"   模板占位符总数：{total_fields}")
    print(f"   已填写字段：{filled_count} 项")
    
    # 严格检查未填写字段
    unfilled_from_config = get_unfilled_fields(field_values, placeholder_groups)
    
    if unfilled_from_config:
        print(f"\n❌ 错误：仍有 {len(unfilled_from_config)} 个字段【必须填写】！")
        print(f"   所有字段都是必填的，除非用户明确说'暂时不填'。")
        print(f"\n   未填写字段列表：")
        for i, f in enumerate(unfilled_from_config[:30]):
            marker = "☐" if f.startswith("☐") else "📝"
            print(f"   {marker} {f}")
        if len(unfilled_from_config) > 30:
            print(f"   ... 及其他 {len(unfilled_from_config) - 30} 项")
        
        if args.check:
            print(f"\n❌ 检查完成：合同未完成填写，请继续向用户询问以上字段。")
            sys.exit(1)
        
        if not args.force:
            print(f"\n❌ 无法生成合同：请先向用户询问并填写以上所有字段。")
            print(f"   提示：继续多轮对话直到所有字段填写完毕。")
            print(f"   如用户明确要求强制生成，请使用 --force 参数。")
            sys.exit(1)
        else:
            print(f"\n⚠️  用户明确要求强制生成，未填字段将留空...")
    
    if args.check:
        print(f"\n✅ 检查完成：所有 {total_fields} 个字段已填写！")
        return
    
    # Fill template
    print(f"\n📝 正在填充模板...")
    unfilled = fill_docx_template(str(template_path), field_values, args.output)
    
    print(f"\n✅ 合同生成成功！")
    print(f"   输出文件：{args.output}")
    
    if unfilled:
        print(f"\n⚠️  模板中有 {len(unfilled)} 个占位符未匹配到值（已清空）：")
        for f in unfilled[:20]:
            print(f"   - {f}")
        if len(unfilled) > 20:
            print(f"   ... 及其他 {len(unfilled) - 20} 项")
    else:
        print(f"\n🎉 所有 {total_fields} 个占位符已完整填写！")


if __name__ == "__main__":
    main()
