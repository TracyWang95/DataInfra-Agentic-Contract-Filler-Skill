# -*- coding: utf-8 -*-
"""
Initialize contract state by parsing a DOCX template for {{placeholder}} markers.
Supports multiple contract types with auto-routing.

Usage:
    # Auto-detect type from user intent:
    python init_contract.py --intent "帮我填数据提供合同" --state "./contract_state.json"
    
    # Specify type explicitly:
    python init_contract.py --type tigong --state "./contract_state.json"
    
    # Use custom template:
    python init_contract.py --type tigong --template "path/to/custom.docx" --state "./state.json"
"""

import os, sys
# Fix Windows console encoding
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

import argparse
import json
import re
from pathlib import Path

# Add parent directory to path for imports
SKILL_DIR = Path(__file__).parent.parent
sys.path.insert(0, str(SKILL_DIR))

from contracts.router import detect_contract_type, CONTRACT_TYPES, list_all_types
from contracts.base_config import get_config, get_progress, get_all_fields


def extract_placeholders_from_docx(template_path: str) -> list[str]:
    """Extract all {{placeholder}} names from a DOCX template."""
    try:
        from docx import Document
    except ImportError:
        print("ERROR: python-docx is required. Install with: pip install python-docx")
        sys.exit(1)
    
    doc = Document(template_path)
    placeholders = set()
    
    # Extract from paragraphs
    for para in doc.paragraphs:
        placeholders.update(re.findall(r'\{\{(.+?)\}\}', para.text))
    
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    placeholders.update(re.findall(r'\{\{(.+?)\}\}', para.text))
    
    return sorted(placeholders)


def init_state(contract_type: str, state_path: str, template_path: str = None):
    """Initialize the contract state file for a specific contract type."""
    
    # Load config for this contract type
    config = get_config(contract_type)
    
    # Use default template if not specified
    if template_path is None:
        template_path = config["TEMPLATE_PATH"]
    
    if not Path(template_path).exists():
        print(f"ERROR: Template file not found: {template_path}")
        sys.exit(1)
    
    # Extract placeholders from template
    placeholders = extract_placeholders_from_docx(str(template_path))
    
    # Classify placeholders
    checkbox_fields = [p for p in placeholders if p.startswith("☐")]
    text_fields = [p for p in placeholders if not p.startswith("☐")]
    
    # Build grouped view from config
    placeholder_groups = config["PLACEHOLDER_GROUPS"]
    grouped = {}
    ungrouped = []
    
    for p in placeholders:
        found = False
        for group_name, group_info in placeholder_groups.items():
            if p in group_info["fields"]:
                if group_name not in grouped:
                    grouped[group_name] = {
                        "description": group_info["description"],
                        "priority": group_info["priority"],
                        "ask": group_info["ask"],
                        "fields": [],
                    }
                grouped[group_name]["fields"].append(p)
                found = True
                break
        if not found:
            ungrouped.append(p)
    
    state = {
        "contract_type": contract_type,
        "contract_name": config["CONTRACT_NAME"],
        "contract_code": config["CONTRACT_CODE"],
        "template_path": str(Path(template_path).resolve()),
        "total_placeholders": len(placeholders),
        "checkbox_count": len(checkbox_fields),
        "text_count": len(text_fields),
        "all_placeholders": placeholders,
        "groups": grouped,
        "ungrouped": ungrouped,
        "field_values": {},  # Will be filled during conversation
    }
    
    # Save state
    state_file = Path(state_path)
    state_file.parent.mkdir(parents=True, exist_ok=True)
    with open(state_file, "w", encoding="utf-8") as f:
        json.dump(state, f, ensure_ascii=False, indent=2)
    
    # Print summary
    type_info = CONTRACT_TYPES[contract_type]
    print(f"✅ 合同初始化完成！")
    print(f"   合同类型：{config['CONTRACT_NAME']}（{config['CONTRACT_CODE']}）")
    print(f"   模板文件：{template_path}")
    print(f"   状态文件：{state_path}")
    print(f"   占位符总数：{len(placeholders)}（文本 {len(text_fields)} + 复选框 {len(checkbox_fields)}）")
    print(f"   分组数量：{len(grouped)}")
    if ungrouped:
        print(f"   未分组字段：{len(ungrouped)} 个")
    print(f"\n📋 分组概览：")
    for group_name in sorted(grouped.keys(), key=lambda g: grouped[g]["priority"]):
        info = grouped[group_name]
        print(f"   [{info['priority']:2d}] {group_name}（{len(info['fields'])} 项）- {info['description']}")


def main():
    parser = argparse.ArgumentParser(description="初始化合同填写状态（支持多种合同类型）")
    parser.add_argument("--type", choices=list(CONTRACT_TYPES.keys()),
                        help="合同类型: tigong/weituo/ronghe/zhongjie")
    parser.add_argument("--intent", help="用户意图描述（用于自动识别合同类型）")
    parser.add_argument("--template", help="自定义模板文件路径（可选，默认使用内置模板）")
    parser.add_argument("--state", help="状态文件保存路径（JSON）")
    parser.add_argument("--list", action="store_true", help="列出所有支持的合同类型")
    args = parser.parse_args()
    
    if args.list:
        print(list_all_types())
        return
    
    # Determine contract type
    contract_type = args.type
    if not contract_type and args.intent:
        contract_type = detect_contract_type(args.intent)
        if contract_type:
            print(f"🔍 识别到合同类型：{CONTRACT_TYPES[contract_type]['name']}")
        else:
            print("❌ 无法识别合同类型。请使用 --type 明确指定，或使用 --list 查看支持的类型。")
            print(list_all_types())
            sys.exit(1)
    
    if not contract_type:
        print("ERROR: 请指定 --type 或 --intent 参数")
        parser.print_help()
        sys.exit(1)
    
    if not args.state:
        print("ERROR: 请指定 --state 参数")
        sys.exit(1)
    
    init_state(contract_type, args.state, args.template)


if __name__ == "__main__":
    main()
